"""AI 面试接口路由。

支持大模型评分；未配置 API Key 时使用本地可解释评分规则兜底。
"""

import json
import re
from datetime import datetime
from io import BytesIO
from typing import Any
from urllib.parse import quote

from fastapi import APIRouter, Depends, HTTPException, Query, status
from fastapi.responses import StreamingResponse
from openai import OpenAI
from sqlalchemy.orm import Session

from app.api.auth import get_current_user
from app.config import settings
from app.database import get_db
from app.models.conversation import Conversation
from app.models.interview import InterviewSession, InterviewTurn
from app.models.user import User
from app.schemas.common import APIResponse, PaginatedResponse
from app.schemas.interview import (
    InterviewAnswerRequest,
    InterviewReportResponse,
    InterviewSessionCreateRequest,
    InterviewSessionDetailResponse,
    InterviewSessionResponse,
)

router = APIRouter(prefix="/api/v1/interview", tags=["AI面试"])

DIMENSIONS = {
    "岗位相关性": 18,
    "技术深度": 24,
    "逻辑结构": 18,
    "案例与结果": 18,
    "表达沟通": 12,
    "时间控制": 10,
}

QUESTION_TEMPLATES = {
    "technical": [
        "请介绍你在{position}方向最有代表性的项目，以及你承担的核心职责。",
        "在{position}工作中遇到复杂问题时，你通常如何定位原因并验证解决方案？",
        "请结合实例说明你如何保证交付质量，并处理性能、稳定性或可维护性问题。",
        "如果需要你从零设计一个与{position}相关的功能，你会如何拆解需求和制定方案？",
        "请谈谈你近期学习的一项与{position}相关的新技术，以及它适合解决什么问题。",
    ],
    "behavioral": [
        "请介绍一次你主动推动团队目标达成的经历。",
        "请描述一次你与团队成员意见不一致的情况，以及你如何处理。",
        "请分享一次你面对紧迫期限时安排优先级的经历。",
        "请说明一次失败或失误给你带来的经验。",
        "你为什么希望从事{position}相关工作，未来的成长目标是什么？",
    ],
    "comprehensive": [
        "请做一个简短的自我介绍，并说明你与{position}岗位的匹配点。",
        "请介绍一个最能体现你解决问题能力的项目经历。",
        "面对不熟悉的任务时，你会如何快速学习并交付结果？",
        "请描述你如何与团队协作并保证信息同步。",
        "你对{position}岗位的核心能力有哪些理解？",
    ],
}


def _llm_enabled() -> bool:
    api_key = (settings.LLM_API_KEY or "").strip()
    return bool(api_key and api_key != "your-api-key-here")


def _call_llm_json(system_prompt: str, user_prompt: str, max_tokens: int = 1600) -> dict[str, Any] | None:
    if not _llm_enabled():
        return None

    try:
        client = OpenAI(
            api_key=settings.LLM_API_KEY,
            base_url=settings.LLM_BASE_URL.rstrip("/"),
            timeout=settings.LLM_TIMEOUT,
        )
        response = client.chat.completions.create(
            model=settings.LLM_MODEL_NAME,
            messages=[
                {"role": "system", "content": system_prompt},
                {"role": "user", "content": user_prompt},
            ],
            temperature=0.0,
            max_tokens=min(settings.LLM_MAX_TOKENS, max_tokens),
            response_format={"type": "json_object"},
        )
        content = response.choices[0].message.content or "{}"
        return json.loads(content)
    except Exception as exc:
        print(f"[AI面试] 大模型调用失败，使用本地规则降级: {exc}")
        return None


def _get_session(db: Session, session_id: int, user_id: int) -> InterviewSession:
    session = db.query(InterviewSession).filter(
        InterviewSession.id == session_id,
        InterviewSession.user_id == user_id,
    ).first()
    if not session:
        raise HTTPException(status_code=status.HTTP_404_NOT_FOUND, detail="面试会话不存在")
    return session


def _generate_question(session: InterviewSession, question_index: int) -> str:
    templates = QUESTION_TEMPLATES.get(session.interview_type, QUESTION_TEMPLATES["comprehensive"])
    template = templates[(question_index - 1) % len(templates)]
    difficulty_label = {"easy": "基础", "medium": "进阶", "hard": "深入"}.get(session.difficulty, "进阶")
    return f"第 {question_index} 题（{difficulty_label}）：{template.format(position=session.target_position)}"


def _get_conversation(db: Session, conversation_id: int | None, user_id: int) -> Conversation | None:
    if conversation_id is None:
        return None
    conversation = db.query(Conversation).filter(
        Conversation.id == conversation_id,
        Conversation.user_id == user_id,
    ).first()
    if not conversation:
        raise HTTPException(status_code=status.HTTP_404_NOT_FOUND, detail="对话不存在")
    return conversation


def _create_turn(db: Session, session: InterviewSession, question_index: int) -> InterviewTurn:
    turn = InterviewTurn(
        session_id=session.id,
        question_index=question_index,
        question=_generate_question(session, question_index),
    )
    db.add(turn)
    db.flush()
    return turn


def _score_band(score: float) -> str:
    if score >= 85:
        return "优秀"
    if score >= 75:
        return "良好"
    if score >= 60:
        return "一般"
    return "待提升"


def _clamp(value: float, max_value: int) -> float:
    return round(max(0.0, min(float(value), float(max_value))), 1)


def _score_answer_locally(
    session: InterviewSession,
    turn: InterviewTurn,
    request: InterviewAnswerRequest,
) -> dict[str, Any]:
    answer = (request.answer_text or "").strip()
    lower_answer = answer.lower()
    length = len(answer)

    structure_markers = ("首先", "其次", "最后", "第一", "第二", "背景", "任务", "行动", "结果", "因此", "总结")
    tech_markers = (
        "接口", "组件", "数据库", "缓存", "并发", "性能", "安全", "测试", "部署", "日志", "异常",
        "架构", "算法", "复杂度", "索引", "事务", "状态", "模块", "封装", "优化",
        "api", "vue", "react", "spring", "mysql", "redis", "docker", "http",
    )
    evidence_markers = ("项目", "负责", "实现", "解决", "上线", "提升", "降低", "用户", "业务", "数据", "指标")
    result_patterns = re.findall(r"\d+[%\w]*|一|二|三|四|五|十|百|千|万", answer)

    dimensions = {
        "岗位相关性": 7.0,
        "技术深度": 8.0,
        "逻辑结构": 6.0,
        "案例与结果": 5.0,
        "表达沟通": 6.0,
        "时间控制": 5.0,
    }

    position_keywords = [word for word in re.split(r"[\s/、，,]+", session.target_position) if word]
    if any(keyword and keyword in answer for keyword in position_keywords):
        dimensions["岗位相关性"] += 5
    if any(marker in answer for marker in ("岗位", "业务", "用户", "需求", "交付", "项目")):
        dimensions["岗位相关性"] += 4
    if length >= 80:
        dimensions["岗位相关性"] += 3

    tech_count = sum(marker in answer or marker in lower_answer for marker in tech_markers)
    dimensions["技术深度"] += min(tech_count * 2.8, 13)
    if any(marker in answer for marker in ("原因", "方案", "验证", "测试", "优化", "权衡")):
        dimensions["技术深度"] += 4

    structure_count = sum(marker in answer for marker in structure_markers)
    dimensions["逻辑结构"] += min(structure_count * 2.2, 9)
    if length >= 120:
        dimensions["逻辑结构"] += 3

    evidence_count = sum(marker in answer for marker in evidence_markers)
    dimensions["案例与结果"] += min(evidence_count * 2.0, 8)
    dimensions["案例与结果"] += min(len(result_patterns) * 1.5, 6)

    if 80 <= length <= 500:
        dimensions["表达沟通"] += 5
    elif length >= 40:
        dimensions["表达沟通"] += 3
    if "我" in answer and any(marker in answer for marker in ("会", "做", "负责", "推进", "解决")):
        dimensions["表达沟通"] += 2

    duration = request.answer_duration_seconds
    if duration is None:
        dimensions["时间控制"] = 6.0
        missing_time = "未记录回答时长，时间控制维度按中等水平估计。"
    elif 60 <= duration <= 180:
        dimensions["时间控制"] = 10.0
        missing_time = ""
    elif 30 <= duration < 60 or 180 < duration <= 240:
        dimensions["时间控制"] = 8.0
        missing_time = ""
    elif 15 <= duration < 30 or 240 < duration <= 300:
        dimensions["时间控制"] = 5.5
        missing_time = "回答时长偏短或偏长，建议控制在 1 到 3 分钟内。"
    else:
        dimensions["时间控制"] = 3.0
        missing_time = "回答时间控制不理想，可能影响面试官对表达完整度和重点把握的判断。"

    dimensions = {key: _clamp(value, DIMENSIONS[key]) for key, value in dimensions.items()}
    score = round(sum(dimensions.values()), 1)

    evidence = []
    missing_points = []
    if tech_count:
        evidence.append(f"回答提到了 {tech_count} 个技术或工程关键词，体现一定技术关联。")
    else:
        missing_points.append("缺少具体技术点、方案选择或实现细节。")
    if structure_count:
        evidence.append("回答包含结构化表达信号，逻辑组织较清楚。")
    else:
        missing_points.append("回答结构不够清晰，建议按背景、任务、行动、结果展开。")
    if evidence_count or result_patterns:
        evidence.append("回答包含项目/结果相关信息，具备一定案例依据。")
    else:
        missing_points.append("缺少可验证的项目案例、数据结果或个人贡献。")
    if length < 60:
        missing_points.append("回答篇幅偏短，论证不够充分。")
    if request.answer_duration_seconds is not None:
        evidence.append(f"本题回答耗时 {request.answer_duration_seconds} 秒，已纳入时间控制评分。")
    if missing_time:
        missing_points.append(missing_time)

    if not evidence:
        evidence.append("回答完成了基本回应，但可支撑评分的具体依据较少。")
    if not missing_points:
        missing_points.append("可以进一步补充量化结果和技术权衡，提升说服力。")

    feedback = f"{_score_band(score)}：本题总分 {score}/100，主要依据为岗位相关性、技术深度、逻辑结构、案例结果、表达沟通和时间控制六个维度。"
    suggestion = missing_points[0]
    return {
        "score": score,
        "dimensions": dimensions,
        "evidence": evidence[:4],
        "missing_points": missing_points[:4],
        "feedback": feedback,
        "suggestion": suggestion,
        "source": "local",
    }


def _score_answer_with_llm(
    session: InterviewSession,
    turn: InterviewTurn,
    request: InterviewAnswerRequest,
) -> dict[str, Any] | None:
    answer = (request.answer_text or "").strip()
    if not answer:
        return None

    system_prompt = """你是一位严谨的中文技术面试官。请按固定评分量表评估候选人回答，只输出 JSON。
必须遵守稳定评分原则：同一个岗位、问题、回答和时长，应给出相同分数；只依据候选人实际回答，不脑补经历；不要因为鼓励候选人而加分。
每个维度先判断档位再给分：空泛/未回答为 0%-35%，有方向但缺少细节为 36%-59%，基本完整为 60%-74%，具体且有方法/案例为 75%-89%，优秀且有量化结果/权衡/验证为 90%-100%。
评分量表总分 100：
1. 岗位相关性 18 分：回答是否贴合目标岗位和问题。
2. 技术深度 24 分：是否包含准确技术点、实现思路、原因分析、取舍或验证方法。
3. 逻辑结构 18 分：是否结构清楚，是否有背景、任务、行动、结果或类似组织方式。
4. 案例与结果 18 分：是否有具体项目、个人贡献、量化结果或可验证事实。
5. 表达沟通 12 分：语言是否清晰、具体、专业。
6. 时间控制 10 分：结合回答时长判断节奏，通常 60 到 180 秒较合适，过短说明展开不足，过长说明重点不够集中。
JSON 字段：
score: 总分数字；
dimensions: 对象，键必须是 岗位相关性/技术深度/逻辑结构/案例与结果/表达沟通/时间控制，值是对应得分；
evidence: 2 到 4 条评分依据；
missing_points: 2 到 4 条扣分原因；
feedback: 一段综合评价；
suggestion: 一条最优先改进建议。"""
    user_prompt = f"""目标岗位：{session.target_position}
面试类型：{session.interview_type}
难度：{session.difficulty}
问题：{turn.question}
回答：{answer}
回答时长秒数：{request.answer_duration_seconds or "未提供"}"""

    data = _call_llm_json(system_prompt, user_prompt)
    if not data:
        return None

    dimensions = data.get("dimensions")
    if not isinstance(dimensions, dict):
        return None
    normalized_dimensions = {
        key: _clamp(dimensions.get(key, 0), max_value)
        for key, max_value in DIMENSIONS.items()
    }
    score = round(sum(normalized_dimensions.values()), 1)
    evidence = [str(item).strip() for item in data.get("evidence", []) if str(item).strip()]
    missing_points = [str(item).strip() for item in data.get("missing_points", []) if str(item).strip()]
    feedback = str(data.get("feedback") or "").strip()
    suggestion = str(data.get("suggestion") or "").strip()
    if not feedback or not suggestion:
        return None

    return {
        "score": score,
        "dimensions": normalized_dimensions,
        "evidence": evidence[:4] or ["大模型根据评分量表完成了综合判断。"],
        "missing_points": missing_points[:4] or ["建议补充更具体的项目细节和结果数据。"],
        "feedback": feedback,
        "suggestion": suggestion,
        "source": "llm",
    }


def _score_answer(session: InterviewSession, turn: InterviewTurn, request: InterviewAnswerRequest) -> dict[str, Any]:
    local_result = _score_answer_locally(session, turn, request)
    llm_result = _score_answer_with_llm(session, turn, request)
    if not llm_result:
        return local_result
    return {
        **local_result,
        "evidence": llm_result.get("evidence") or local_result["evidence"],
        "missing_points": llm_result.get("missing_points") or local_result["missing_points"],
        "feedback": llm_result.get("feedback") or local_result["feedback"],
        "suggestion": llm_result.get("suggestion") or local_result["suggestion"],
        "source": "local_score_llm_feedback",
    }


def _analyze_turn_for_report(session: InterviewSession, turn: InterviewTurn) -> dict[str, Any]:
    fake_request = InterviewAnswerRequest(
        answer_text=turn.answer_text or "",
        answer_duration_seconds=turn.answer_duration_seconds,
        answer_audio_url=turn.answer_audio_url,
    )
    detail = _score_answer_locally(session, turn, fake_request)
    detail["score"] = float(turn.score or detail["score"])
    detail["feedback"] = turn.feedback or detail["feedback"]
    detail["suggestion"] = turn.suggestion or detail["suggestion"]
    return detail


def _average_dimensions(turn_details: list[dict[str, Any]]) -> dict[str, float]:
    if not turn_details:
        return {key: 0 for key in DIMENSIONS}
    totals = {key: 0.0 for key in DIMENSIONS}
    for detail in turn_details:
        for key in DIMENSIONS:
            totals[key] += float(detail.get("dimensions", {}).get(key, 0))
    return {key: round(value / len(turn_details), 1) for key, value in totals.items()}


def _build_local_report(session: InterviewSession) -> dict[str, Any]:
    answered_turns = [turn for turn in session.turns if turn.answered_at is not None]
    turn_details = [_analyze_turn_for_report(session, turn) for turn in answered_turns]
    scores = [float(turn.score or 0) for turn in answered_turns]
    answered_count = len(scores)
    total_score = round(sum(scores) / answered_count, 1) if answered_count else 0.0
    dimension_scores = _average_dimensions(turn_details)

    low_dimensions = sorted(dimension_scores.items(), key=lambda item: item[1] / DIMENSIONS[item[0]])[:2]
    high_dimensions = sorted(dimension_scores.items(), key=lambda item: item[1] / DIMENSIONS[item[0]], reverse=True)[:2]
    strengths = [f"{name}表现相对较好，平均得分 {score}/{DIMENSIONS[name]}。" for name, score in high_dimensions if score > 0]
    weaknesses = [f"{name}仍需加强，平均得分 {score}/{DIMENSIONS[name]}。" for name, score in low_dimensions if score > 0]
    if not strengths:
        strengths = ["已完成全部题目回答，具备进一步复盘和提升的基础。"]
    if not weaknesses:
        weaknesses = ["建议继续提升回答的具体案例、技术细节和结果数据。"]

    suggestions = list(dict.fromkeys(
        detail["suggestion"] for detail in turn_details if detail.get("suggestion")
    ))
    if len(suggestions) < 3:
        suggestions.extend([
            "每道题尽量说明具体场景、个人行动和最终结果。",
            "技术题回答时补充方案选择、验证方法和异常处理。",
            "准备 2 到 3 个可复用项目案例，用数据说明贡献。",
        ])
    suggestions = list(dict.fromkeys(suggestions))[:5]

    action_plan = [
        "第 1 步：整理一个最熟悉的项目，用背景、任务、行动、结果四段话讲清楚。",
        "第 2 步：为每个项目补充技术难点、方案对比、测试验证和量化结果。",
        "第 3 步：将常见问题录音练习，控制回答在 1 到 3 分钟内。",
        "第 4 步：针对低分维度重新回答本次题目，并对照评分依据自查。",
    ]

    if total_score >= 85:
        summary = "整体表现优秀，回答较完整，能体现较好的岗位匹配度和表达能力。后续重点是继续补充量化结果和技术取舍，让优势更有说服力。"
    elif total_score >= 75:
        summary = "整体表现良好，多数回答能够回应问题，但部分题目还可以补充更具体的项目细节、技术依据和结果数据。"
    elif total_score >= 60:
        summary = "整体表现一般，回答覆盖了基本方向，但技术深度、案例依据或结构化表达仍不够稳定。建议按评分维度逐项补强。"
    else:
        summary = "整体表现有较大提升空间，回答中缺少足够的具体事实和结构化论证。建议先准备标准项目案例，再进行模拟面试训练。"

    return {
        "session_id": session.id,
        "total_score": total_score,
        "status": session.status,
        "summary": summary,
        "score_basis": f"总分按已回答题目平均分计算：本次已回答 {answered_count} 道题，每题先按六个维度评分（岗位相关性18分、技术深度24分、逻辑结构18分、案例与结果18分、表达沟通12分、时间控制10分），单题满分100分；最终总分 = 已回答题目分数之和 / 已回答题数。未回答题目不参与平均分，时间控制会结合每题回答秒数，通常60到180秒较合适。",
        "dimension_scores": dimension_scores,
        "turn_performance": [
            {
                "question_index": turn.question_index,
                "question": turn.question,
                "answer": turn.answer_text or "",
                "answer_duration_seconds": turn.answer_duration_seconds,
                "score": float(turn.score or 0),
                "dimensions": detail["dimensions"],
                "evidence": detail["evidence"],
                "missing_points": detail["missing_points"],
                "feedback": turn.feedback or detail["feedback"],
                "suggestion": turn.suggestion or detail["suggestion"],
            }
            for turn, detail in zip(answered_turns, turn_details)
        ],
        "strengths": strengths[:4],
        "weaknesses": weaknesses[:4],
        "suggestions": suggestions,
        "action_plan": action_plan,
        "generated_at": datetime.now(),
    }


def _enhance_report_with_llm(report: dict[str, Any], session: InterviewSession) -> dict[str, Any]:
    if not report["turn_performance"]:
        return report

    system_prompt = """你是一位中文面试辅导专家。请基于评分表和逐题记录生成详细面试报告，只输出 JSON。
必须保留客观依据，不要编造候选人未提到的经历。
JSON 字段：
summary: 3 到 5 句话整体总结；
strengths: 3 到 5 条优势；
weaknesses: 3 到 5 条不足；
suggestions: 4 到 6 条改进建议；
action_plan: 4 到 6 条后续训练计划。"""
    user_prompt = json.dumps(
        {
            "target_position": session.target_position,
            "interview_type": session.interview_type,
            "difficulty": session.difficulty,
            "score_basis": report["score_basis"],
            "total_score": report["total_score"],
            "dimension_scores": report["dimension_scores"],
            "turn_performance": report["turn_performance"],
        },
        ensure_ascii=False,
    )
    data = _call_llm_json(system_prompt, user_prompt, max_tokens=2200)
    if not data:
        return report

    for key in ["summary", "strengths", "weaknesses", "suggestions", "action_plan"]:
        value = data.get(key)
        if key == "summary" and isinstance(value, str) and value.strip():
            report[key] = value.strip()
        elif isinstance(value, list) and value:
            report[key] = [str(item).strip() for item in value if str(item).strip()]
    return report


def _build_report(session: InterviewSession) -> dict[str, Any]:
    return _enhance_report_with_llm(_build_local_report(session), session)


def _save_report(db: Session, session: InterviewSession) -> InterviewReportResponse:
    report = _build_report(session)
    session.total_score = report["total_score"]
    session.report = {**report, "generated_at": report["generated_at"].isoformat()}
    db.commit()
    db.refresh(session)
    return InterviewReportResponse.model_validate(report)


def _write_docx_report(session: InterviewSession, report: dict[str, Any]) -> BytesIO:
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.shared import Pt

    def apply_run_font(run, size: int = 12, bold: bool = False) -> None:
        run.bold = bold
        run.font.name = "宋体"
        run._element.rPr.rFonts.set(qn("w:ascii"), "宋体")
        run._element.rPr.rFonts.set(qn("w:hAnsi"), "宋体")
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
        run.font.size = Pt(size)

    def add_plain_paragraph(
        text: Any,
        size: int = 12,
        bold: bool = False,
        alignment: WD_ALIGN_PARAGRAPH | None = None,
    ):
        paragraph = document.add_paragraph()
        if alignment is not None:
            paragraph.alignment = alignment
        run = paragraph.add_run(str(text))
        apply_run_font(run, size=size, bold=bold)
        return paragraph

    def set_cell_text(cell, value: Any) -> None:
        cell.text = str(value)
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                apply_run_font(run)

    document = Document()
    for style in document.styles:
        if hasattr(style, "font"):
            style.font.name = "宋体"
            style._element.rPr.rFonts.set(qn("w:ascii"), "宋体")
            style._element.rPr.rFonts.set(qn("w:hAnsi"), "宋体")
            style._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
            style.font.size = Pt(12)

    add_plain_paragraph(
        f"{session.title} - 面试评估报告",
        size=20,
        bold=True,
        alignment=WD_ALIGN_PARAGRAPH.CENTER,
    )
    add_plain_paragraph(f"目标岗位：{session.target_position}")
    add_plain_paragraph(f"面试类型：{session.interview_type}    难度：{session.difficulty}")
    add_plain_paragraph(f"综合评分：{report.get('total_score', 0)}/100")
    add_plain_paragraph(f"生成时间：{report.get('generated_at')}")

    add_plain_paragraph("一、评分依据", bold=True)
    add_plain_paragraph(report.get("score_basis") or "")

    add_plain_paragraph("二、总体评价", bold=True)
    add_plain_paragraph(report.get("summary") or "")

    add_plain_paragraph("三、维度得分", bold=True)
    table = document.add_table(rows=1, cols=3)
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    set_cell_text(hdr[0], "评分维度")
    set_cell_text(hdr[1], "得分")
    set_cell_text(hdr[2], "满分")
    for name, score in (report.get("dimension_scores") or {}).items():
        row = table.add_row().cells
        set_cell_text(row[0], name)
        set_cell_text(row[1], score)
        set_cell_text(row[2], DIMENSIONS.get(name, ""))

    add_plain_paragraph("四、逐题分析", bold=True)
    for turn in report.get("turn_performance") or []:
        add_plain_paragraph(f"第 {turn.get('question_index')} 题：{turn.get('score')}/100", bold=True)
        add_plain_paragraph(f"题目：{turn.get('question')}")
        add_plain_paragraph(f"回答：{turn.get('answer') or '未记录'}")
        duration = turn.get("answer_duration_seconds")
        add_plain_paragraph(f"回答时长：{duration if duration is not None else '未记录'} 秒")
        add_plain_paragraph(f"评价：{turn.get('feedback')}")
        add_plain_paragraph("评分依据：")
        for item in turn.get("evidence") or []:
            add_plain_paragraph(f"- {item}")
        add_plain_paragraph("扣分点：")
        for item in turn.get("missing_points") or []:
            add_plain_paragraph(f"- {item}")
        add_plain_paragraph(f"改进建议：{turn.get('suggestion')}")

    add_plain_paragraph("五、优势与不足", bold=True)
    add_plain_paragraph("优势：")
    for item in report.get("strengths") or []:
        add_plain_paragraph(f"- {item}")
    add_plain_paragraph("不足：")
    for item in report.get("weaknesses") or []:
        add_plain_paragraph(f"- {item}")

    add_plain_paragraph("六、改进建议与训练计划", bold=True)
    for item in report.get("suggestions") or []:
        add_plain_paragraph(f"- {item}")
    add_plain_paragraph("训练计划：")
    for index, item in enumerate(report.get("action_plan") or [], start=1):
        add_plain_paragraph(f"{index}. {item}")

    output = BytesIO()
    document.save(output)
    output.seek(0)
    return output

@router.post("/sessions", response_model=APIResponse[InterviewSessionResponse], summary="创建面试会话")
async def create_session(
    request: InterviewSessionCreateRequest,
    current_user: User = Depends(get_current_user),
    db: Session = Depends(get_db),
):
    _get_conversation(db, request.conversation_id, current_user.id)
    session = InterviewSession(
        user_id=current_user.id,
        conversation_id=request.conversation_id,
        title=request.title or f"{request.target_position} AI面试",
        target_position=request.target_position,
        interview_type=request.interview_type,
        difficulty=request.difficulty,
        question_count=request.question_count,
        answer_mode=request.answer_mode,
    )
    db.add(session)
    db.commit()
    db.refresh(session)
    return APIResponse(data=InterviewSessionResponse.model_validate(session), message="面试会话创建成功")


@router.get("/sessions", response_model=APIResponse[PaginatedResponse[InterviewSessionResponse]], summary="获取面试会话列表")
async def list_sessions(
    page: int = Query(default=1, ge=1),
    page_size: int = Query(default=20, ge=1, le=100),
    conversation_id: int | None = Query(default=None, ge=1),
    current_user: User = Depends(get_current_user),
    db: Session = Depends(get_db),
):
    _get_conversation(db, conversation_id, current_user.id)
    query = db.query(InterviewSession).filter(InterviewSession.user_id == current_user.id)
    if conversation_id is not None:
        query = query.filter(InterviewSession.conversation_id == conversation_id)
    total = query.count()
    sessions = (
        query.order_by(InterviewSession.created_at.desc(), InterviewSession.id.desc())
        .offset((page - 1) * page_size)
        .limit(page_size)
        .all()
    )
    return APIResponse(
        data=PaginatedResponse(
            data=[InterviewSessionResponse.model_validate(item) for item in sessions],
            total=total,
            page=page,
            page_size=page_size,
        )
    )


@router.get("/sessions/{session_id}", response_model=APIResponse[InterviewSessionDetailResponse], summary="获取面试会话详情")
async def get_session_detail(
    session_id: int,
    current_user: User = Depends(get_current_user),
    db: Session = Depends(get_db),
):
    session = _get_session(db, session_id, current_user.id)
    return APIResponse(data=InterviewSessionDetailResponse.model_validate(session))


@router.post("/sessions/{session_id}/start", response_model=APIResponse[InterviewSessionDetailResponse], summary="开始面试")
async def start_session(
    session_id: int,
    current_user: User = Depends(get_current_user),
    db: Session = Depends(get_db),
):
    session = _get_session(db, session_id, current_user.id)
    if session.status == "finished":
        raise HTTPException(status_code=status.HTTP_400_BAD_REQUEST, detail="面试已结束")

    if session.status == "pending":
        session.status = "in_progress"
        session.started_at = datetime.now()

    if not session.turns:
        _create_turn(db, session, 1)

    db.commit()
    db.refresh(session)
    return APIResponse(data=InterviewSessionDetailResponse.model_validate(session), message="面试已开始")


@router.post("/sessions/{session_id}/answer", response_model=APIResponse[InterviewSessionDetailResponse], summary="提交当前问题回答")
async def answer_session(
    session_id: int,
    request: InterviewAnswerRequest,
    current_user: User = Depends(get_current_user),
    db: Session = Depends(get_db),
):
    session = _get_session(db, session_id, current_user.id)
    if session.status != "in_progress":
        raise HTTPException(status_code=status.HTTP_400_BAD_REQUEST, detail="面试尚未开始或已经结束")

    turn = next((item for item in session.turns if item.answered_at is None), None)
    if not turn:
        raise HTTPException(status_code=status.HTTP_400_BAD_REQUEST, detail="当前没有待回答的问题")

    result = _score_answer(session, turn, request)
    turn.answer_text = request.answer_text.strip() if request.answer_text else None
    turn.answer_audio_url = request.answer_audio_url
    turn.answer_duration_seconds = request.answer_duration_seconds
    turn.answered_at = datetime.now()
    turn.score = result["score"]
    turn.feedback = result["feedback"]
    turn.suggestion = result["suggestion"]

    if turn.question_index < session.question_count:
        _create_turn(db, session, turn.question_index + 1)

    db.commit()
    db.refresh(session)
    return APIResponse(data=InterviewSessionDetailResponse.model_validate(session), message="回答提交成功")


@router.post("/sessions/{session_id}/finish", response_model=APIResponse[InterviewReportResponse], summary="结束面试")
async def finish_session(
    session_id: int,
    current_user: User = Depends(get_current_user),
    db: Session = Depends(get_db),
):
    session = _get_session(db, session_id, current_user.id)
    if session.status == "pending":
        raise HTTPException(status_code=status.HTTP_400_BAD_REQUEST, detail="面试尚未开始")

    if session.status != "finished":
        session.status = "finished"
        session.ended_at = datetime.now()

    report = _save_report(db, session)
    return APIResponse(data=report, message="面试已结束")


@router.get("/sessions/{session_id}/report", response_model=APIResponse[InterviewReportResponse], summary="获取面试报告")
async def get_session_report(
    session_id: int,
    current_user: User = Depends(get_current_user),
    db: Session = Depends(get_db),
):
    session = _get_session(db, session_id, current_user.id)
    report = _save_report(db, session)
    return APIResponse(data=report)


@router.get("/sessions/{session_id}/report.docx", summary="下载 Word 面试报告")
async def download_session_report_docx(
    session_id: int,
    current_user: User = Depends(get_current_user),
    db: Session = Depends(get_db),
):
    session = _get_session(db, session_id, current_user.id)
    report = _save_report(db, session).model_dump()
    output = _write_docx_report(session, report)
    filename = quote(f"{session.title}-面试报告.docx")
    return StreamingResponse(
        output,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": f"attachment; filename*=UTF-8''{filename}"},
    )
